from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SOURCE = Path(
    r"C:\Users\heheg\Downloads\Future HCI conference paper structure _ updated.docx"
)
OUTPUT = Path(
    r"C:\Users\heheg\OneDrive\Desktop\FAA map project"
    r"\Future HCI conference paper structure - presentation highlights.docx"
)


YELLOW = {
    10: [
        "This paper presents a vision for the future of ATC/ATM training through immersive, adaptive, and collaborative learning ecosystems.",
        "Building upon this foundation, we envision future ATC/ATM training ecosystems centered around three complementary educational modalities: self-paced learning, AI-supported scenario-based practice, and instructor-led collaborative learning.",
    ],
    15: [
        "Air traffic control and management (ATC/ATM) remains a safety-critical occupation with exceptionally high training demands.",
        "The Federal Aviation Administration (FAA) recently reached a 27-year low in controller staffing, with nearly one-third of controllers projected to become retirement-eligible by 2030 and substantial workforce losses anticipated throughout the 2021–2030 period [6].",
        "As a result, trainees must rapidly internalize dense spatial and procedural information within instructional environments that still rely heavily on static memorization-based learning approaches [17].",
        "These challenges have intensified the need for scalable, human-centered training systems that improve engagement, spatial understanding, knowledge retention, and workforce readiness.",
    ],
    17: [
        "Current research and applied training frameworks show limited integration of XR technologies for scalable instruction, collaborative training, and adaptive educational support.",
    ],
    18: [
        "We therefore treat ATC/ATM not only as an application area but as a high-complexity testbed for a general HCI question: how can immersive, adaptive, and collaborative learning ecosystems be designed to support expertise development in safety-critical, spatially demanding professions?",
    ],
    22: [
        "In contrast, immersive 3D visualization presents these relationships directly, reducing the burden of mental reconstruction while making spatial context easier to interpret.",
    ],
    29: [
        "Although XR offers significant promise for aviation training, safety-critical immersive systems must be designed around human needs, cognitive limitations, and operational demands rather than technical novelty alone.",
        "Within ATC/ATM training environments, higher realism is not automatically beneficial.",
    ],
    35: [
        "By supporting browser-based deployment and synchronized multi-user environments, WebXR expands access to immersive training while preserving embodied VR interaction across devices.",
        "These capabilities support self-paced learning, instructor-guided collaboration, and geographically distributed participation, making WebXR a practical foundation for future ATC/ATM learning ecosystems.",
    ],
    37: [
        "These discussions consistently highlighted VFR map memorization as a major obstacle during early training and identified immersive VR as a promising modality for supporting spatial understanding, engagement, and experiential learning [22].",
    ],
    38: [
        "Successive usability evaluations with novice users and CPCs demonstrated strong usability (SUS = 73.2–85) [1,9], perceived training value, and opportunities for enhanced realism, instructional guidance, and instructor-facing functionality [17,22].",
        "These recommendations directly informed the current generation of the platform, which migrated to a WebXR framework to support scalability, dynamic FAA data integration, and collaborative learning.",
        "Additional enhancements include controller-based map manipulation, integrated 3D airspace structures, and synchronized interactions across linked sessions, establishing the technical foundation for future instructor-guided learning environments.",
    ],
    40: [
        "The first step is to evaluate whether immersive 3D airspace visualization provides measurable advantages over traditional 2D instructional methods for learning VFR sectional information.",
        "Controlled studies will assess spatial understanding, map memorization, knowledge retention, and learner self-efficacy.",
        "Critically, efficacy measurement is positioned at the earliest stage of the validation pathway rather than deferred until later ecosystem features are built, so that higher-order capabilities are grounded in demonstrated learning benefits rather than assumed ones.",
    ],
    48: [
        "Rather than relying primarily on memorization and procedural instruction, trainees could engage with realistic operational situations that require active problem-solving, decision-making, and application of learned concepts.",
        "AI-assisted scenario generation may support a wide range of training experiences, from routine traffic management operations to emergency events and weather-related disruptions, while allowing instructors to dynamically adjust scenario complexity and conditions to meet instructional objectives.",
    ],
    51: [
        "Although AI and self-paced learning may expand access to training opportunities, instructors will remain a critical component of future learning ecosystems.",
        "Future collaborative environments may support one-to-many instruction within synchronized immersive spaces where instructors demonstrate operational techniques, discuss decision-making strategies, and guide trainees through realistic scenarios.",
    ],
    56: [
        "Although grounded in ATC/ATM training, the contribution of this work is a transferable model rather than an aviation-specific result.",
    ],
    58: [
        "We suggest a domain is a strong candidate when it exhibits: (1) high spatial and procedural complexity that requires learners to reconstruct three-dimensional structure from two-dimensional representations; (2) a safety-critical context in which real-world practice is risky, costly, or difficult to reproduce; (3) reliance on situational awareness and information prioritization across a complex operational display; (4) expertise developed through instructor-led apprenticeship combined with repeated rehearsal and self-efficacy development; and (5) a need for scenario variety that is difficult to stage consistently in reality, favoring AI-assisted generation.",
    ],
    62: [
        "Current AI-supported learning research has not established how instructional authority should be distributed among the learner, instructor, and AI system, particularly in collaborative immersive environments.",
    ],
    63: [
        "The longer-term ATC/ATM platform could provide a testbed for examining how different allocations of AI authority affect learner agency, instructor decision-making, and independent learning.",
        "Studies could compare direct AI guidance, instructor-mediated recommendations, learner-requested assistance, and conditions without AI support.",
    ],
    66: [
        "In a VR/tablet pair-learning study, asymmetric and symmetric conditions produced comparable learning outcomes, although the symmetric condition provided stronger presence, immersion, and awareness of the partner’s activity.",
        "However, these systems have received limited classroom evaluation.",
    ],
    67: [
        "Synchronized highlighting, layer states, task instructions, and event overlays would provide shared instructional references across the desktop and immersive interfaces.",
        "This work would address a question that remains insufficiently studied: what information and controls must be shared across asymmetric interfaces for one instructor to guide several immersive learners effectively?",
    ],
    70: [
        "Assessment remains a central challenge in immersive learning research because usability, presence, engagement, and confidence are often reported without establishing whether learners retain knowledge or transfer it beyond VR.",
        "Reviews of safety-relevant training show that few studies evaluate subsequent behavior, long-term outcomes, or transfer to real tasks (Stefan et al. 2023), while experimental evidence demonstrates that stronger presence does not necessarily produce better learning (Makransky et al., 2019).",
        "Comparisons between immersive visualization and paper charts should combine immediate and delayed knowledge tests with transfer tasks using unfamiliar chart content.",
        "Because recorded behavior may reflect interface proficiency rather than aviation knowledge, system telemetry should not be interpreted as competence without external validation.",
    ],
    76: [
        "Modern ATC/ATM operations require trainees to understand dynamic three-dimensional airspace, yet early training continues to rely heavily on static charts and two-dimensional displays, creating a mismatch between how airspace is represented during training and how it must be understood in practice.",
        "Learning ecosystems that combine self-paced study, scenario-based practice, and instructor-led collaboration may expand access to training, strengthen self-efficacy, and support workforce development.",
    ],
}


GREEN = {
    60: [
        "Discussion prompt: Which of these properties are necessary versus merely sufficient for transfer, and what safety-critical domains does the model fail to capture?",
    ],
    64: [
        "Discussion prompt: How can AI-mediated training augment rather than displace human instructors while preserving learner agency in safety-critical domains?",
    ],
    68: [
        "Discussion prompt: What forms of co-presence and shared spatial grounding are required for effective one-to-many instruction in immersive environments?",
    ],
    71: [
        "Discussion prompt: What should count as evidence of learning in immersive ecosystems, and how early must efficacy be demonstrated before building on unvalidated assumptions?",
    ],
    74: [
        "Discussion prompt: How can browser-based immersive infrastructure broaden access to professional training without reproducing existing inequities?",
    ],
}


def collect_ranges(text, excerpts, color):
    ranges = []
    for excerpt in excerpts:
        start = text.find(excerpt)
        if start < 0:
            raise ValueError(f"Excerpt not found: {excerpt}")
        ranges.append((start, start + len(excerpt), color))
    return ranges


def highlight_paragraph(paragraph, ranges):
    text = paragraph.text
    source_runs = [run for run in paragraph.runs if run.text]
    source_rpr = deepcopy(source_runs[0]._r.rPr) if source_runs else None

    boundaries = {0, len(text)}
    for start, end, _ in ranges:
        boundaries.add(start)
        boundaries.add(end)
    boundaries = sorted(boundaries)

    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    for start, end in zip(boundaries, boundaries[1:]):
        if start == end:
            continue
        run = paragraph.add_run(text[start:end])
        if source_rpr is not None:
            run._r.insert(0, deepcopy(source_rpr))
        for range_start, range_end, color in ranges:
            if start >= range_start and end <= range_end:
                run.font.highlight_color = color
                break


def main():
    document = Document(SOURCE)
    selected = {}

    for paragraph_index, excerpts in YELLOW.items():
        selected.setdefault(paragraph_index, []).extend(
            collect_ranges(
                document.paragraphs[paragraph_index].text,
                excerpts,
                WD_COLOR_INDEX.YELLOW,
            )
        )
    for paragraph_index, excerpts in GREEN.items():
        selected.setdefault(paragraph_index, []).extend(
            collect_ranges(
                document.paragraphs[paragraph_index].text,
                excerpts,
                WD_COLOR_INDEX.GRAY_25,
            )
        )

    for paragraph_index, ranges in selected.items():
        highlight_paragraph(document.paragraphs[paragraph_index], ranges)

    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Highlighted excerpts: {sum(len(v) for v in YELLOW.values())}")
    print(f"Highlighted discussion prompts: {sum(len(v) for v in GREEN.values())}")


if __name__ == "__main__":
    main()
